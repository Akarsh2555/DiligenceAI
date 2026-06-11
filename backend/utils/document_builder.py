import os
import re
from pathlib import Path
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_chart(title: str, labels: list, values: list, filename: str, color: str):
    """Generates a beautiful minimalist bar chart and saves it as an image."""
    plt.figure(figsize=(6, 3), facecolor='white')
    # Minimalist style
    ax = plt.gca()
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_color('#E0E0E0')
    ax.tick_params(left=False, bottom=False)
    
    bars = plt.bar(labels, values, color=color, width=0.5, alpha=0.8)
    plt.title(title, pad=20, fontsize=12, fontweight='bold', color='#333333')
    
    # Add values on top of bars
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{height}',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3),  # 3 points vertical offset
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=10, color='#666666')
        
    plt.tight_layout()
    plt.savefig(filename, dpi=300, bbox_inches='tight')
    plt.close()

def apply_beautiful_styles(doc):
    """Applies modern styling to the Document."""
    styles = doc.styles
    
    # Title Style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Helvetica'
    title_font.size = Pt(28)
    title_font.color.rgb = RGBColor(0, 51, 102) # Dark Navy
    
    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Helvetica'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 102, 204) # Vibrant Blue
    
    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Helvetica'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(51, 51, 51) # Dark Gray
    
    # Normal Text
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(80, 80, 80) # Soft Gray

def parse_markdown_to_docx(doc, text: str):
    """Extremely basic markdown parser to add text to the docx."""
    if not text:
        return
        
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold inside bullet
            text_part = line[2:].strip()
            if '**' in text_part:
                parts = text_part.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(text_part)
        else:
            p = doc.add_paragraph()
            # Handle bold inside normal paragraph
            if '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(line)

def generate_beautiful_report(report_data: dict, output_path: str):
    """
    Acts as the core tool to convert LangGraph output into a stunning DOCX file.
    It generates dynamic graphs based on risk vs growth metrics if available.
    """
    doc = Document()
    apply_beautiful_styles(doc)
    
    # --- Title Page ---
    title = doc.add_paragraph("DiligenceAI", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Comprehensive Due Diligence Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    parse_markdown_to_docx(doc, report_data.get('executive_summary', 'No summary available.'))
    
    # Count the number of bullet points (insights) the AI extracted for each section
    def count_insights(text):
        return len(re.findall(r'(?m)^[-*]\s', str(text)))

    risk_count = count_insights(report_data.get('risk_assessment', ''))
    growth_count = count_insights(report_data.get('growth_opportunities', ''))
    legal_count = count_insights(report_data.get('legal_analysis', ''))

    # Generate an actual metric chart based on the AI's findings
    chart_path = "temp_sentiment_chart.png"
    create_chart(
        title="Number of Insights Extracted from Document",
        labels=['Risk Factors', 'Growth Signals', 'Legal Clauses'],
        values=[risk_count, growth_count, legal_count],
        filename=chart_path,
        color='#0066CC'
    )
    doc.add_picture(chart_path, width=Inches(5.0))
    os.remove(chart_path)
    
    doc.add_page_break()
    
    # --- Risk Assessment ---
    doc.add_heading('Risk Assessment', level=1)
    parse_markdown_to_docx(doc, report_data.get('risk_assessment', 'No risk assessment.'))
    
    # Adding a Table for Risk Matrix
    doc.add_heading('Risk Matrix', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Severity'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'Recommendation'
    
    # Add dynamic rows based on actual risk assessment text
    risk_text = str(report_data.get('risk_assessment', ''))
    
    # Simple heuristic to extract risks: find bullet points
    risks = re.findall(r'(?m)^[-*]\s(.*?)$', risk_text)
    
    if not risks:
        row_cells = table.add_row().cells
        row_cells[0].text = 'None'
        row_cells[1].text = 'N/A'
        row_cells[2].text = 'No significant risks identified'
    else:
        for risk in risks[:5]: # Top 5 risks
            row_cells = table.add_row().cells
            
            # Try to guess severity from text
            severity = 'Medium'
            if 'high' in risk.lower() or 'critical' in risk.lower() or 'severe' in risk.lower():
                severity = 'High'
            elif 'low' in risk.lower() or 'minor' in risk.lower():
                severity = 'Low'
                
            row_cells[0].text = severity
            row_cells[1].text = 'Operational/Legal'
            # Remove markdown bold/italic formatting for clean table insertion
            clean_risk = re.sub(r'[*_]', '', risk)
            row_cells[2].text = clean_risk[:100] + ('...' if len(clean_risk) > 100 else '')
    
    doc.add_page_break()
    
    # --- Growth Opportunities ---
    doc.add_heading('Growth Opportunities', level=1)
    parse_markdown_to_docx(doc, report_data.get('growth_opportunities', 'No growth data.'))
    
    doc.add_page_break()
    
    # --- Legal & Regulatory ---
    doc.add_heading('Legal Analysis', level=1)
    parse_markdown_to_docx(doc, report_data.get('legal_analysis', 'No legal data.'))
    
    # Save the document
    doc.save(output_path)
    return output_path
